"""
Document Processing Service for Threat Modeling

This service handles:
- File upload and parsing (PDF, DOCX, TXT)
- Text extraction and processing
- Document storage and management
- Integration with threat models
"""

import os
import io
import uuid
from typing import Dict, List, Optional, Any, BinaryIO
from datetime import datetime
import logging
from pathlib import Path

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process and extract text from various document formats"""
    
    SUPPORTED_FORMATS = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
        'doc': 'application/msword',
        'txt': 'text/plain'
    }
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB limit
    
    def __init__(self):
        """Initialize document processor"""
        self.upload_dir = Path("uploads/documents")
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    async def process_document(self, 
                             file_content: bytes, 
                             filename: str, 
                             content_type: str) -> Dict[str, Any]:
        """
        Process uploaded document and extract text content
        
        Args:
            file_content: Binary content of the file
            filename: Original filename
            content_type: MIME type of the file
        
        Returns:
            Dictionary with processing results
        """
        try:
            # Validate file
            validation_result = self._validate_file(file_content, filename, content_type)
            if not validation_result["valid"]:
                return validation_result
            
            # Determine file format
            file_format = self._get_file_format(filename, content_type)
            if not file_format:
                return {
                    "success": False,
                    "error": f"Unsupported file format. Supported formats: {list(self.SUPPORTED_FORMATS.keys())}"
                }
            
            logger.info(f"📄 Processing {file_format.upper()} document: {filename}")
            
            # Extract text based on format
            text_content = await self._extract_text(file_content, file_format)
            
            if not text_content.strip():
                return {
                    "success": False,
                    "error": "No text content could be extracted from the document"
                }
            
            logger.info(f"📄 Extracted {len(text_content)} characters from {filename}")
            
            # Validate document relevance for threat modeling
            relevance_check = self._validate_document_relevance(text_content, filename)
            
            if not relevance_check['is_relevant']:
                logger.warning(f"📄 Document '{filename}' appears unrelated to threat modeling (confidence: {relevance_check['confidence']}%)")
                return {
                    "success": False,
                    "error": f"Document appears unrelated to threat modeling, security, or system architecture. Please upload technical documentation, architecture diagrams, requirements, or security-related content. (Relevance confidence: {relevance_check['confidence']}%)",
                    "relevance_analysis": relevance_check
                }
            
            logger.info(f"📄 Document relevance validated (score: {relevance_check['relevance_score']}, categories: {relevance_check['found_categories']})")
            
            # Generate document ID and save file
            document_id = str(uuid.uuid4())
            file_path = await self._save_file(file_content, document_id, file_format)
            
            # Analyze document content for threat modeling context
            analysis = self._analyze_document_content(text_content)
            
            result = {
                "success": True,
                "document_id": document_id,
                "filename": filename,
                "file_format": file_format,
                "file_size": len(file_content),
                "text_content": text_content,
                "word_count": len(text_content.split()),
                "character_count": len(text_content),
                "file_path": str(file_path),
                "processed_at": datetime.utcnow().isoformat(),
                "analysis": analysis
            }
            
            logger.info(f"✅ Document processed successfully: {filename} ({len(text_content)} chars)")
            return result
            
        except Exception as e:
            logger.error(f"❌ Error processing document {filename}: {str(e)}")
            return {
                "success": False,
                "error": f"Document processing failed: {str(e)}"
            }
    
    def _validate_file(self, file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """Validate uploaded file"""
        
        # Check file size
        if len(file_content) > self.MAX_FILE_SIZE:
            return {
                "valid": False,
                "success": False,
                "error": f"File too large. Maximum size: {self.MAX_FILE_SIZE // (1024*1024)}MB"
            }
        
        # Check if file is empty
        if len(file_content) == 0:
            return {
                "valid": False,
                "success": False,
                "error": "File is empty"
            }
        
        return {"valid": True}
    
    def _get_file_format(self, filename: str, content_type: str) -> Optional[str]:
        """Determine file format from filename and content type"""
        
        # Check by file extension
        if filename:
            extension = Path(filename).suffix.lower().lstrip('.')
            if extension in self.SUPPORTED_FORMATS:
                return extension
        
        # Check by content type
        for format_ext, mime_type in self.SUPPORTED_FORMATS.items():
            if content_type == mime_type:
                return format_ext
        
        return None
    
    async def _extract_text(self, file_content: bytes, file_format: str) -> str:
        """Extract text content based on file format"""
        
        try:
            if file_format == 'pdf':
                return await self._extract_pdf_text(file_content)
            elif file_format in ['docx', 'doc']:
                return await self._extract_docx_text(file_content)
            elif file_format == 'txt':
                return await self._extract_txt_text(file_content)
            else:
                raise ValueError(f"Unsupported format: {file_format}")
        
        except Exception as e:
            logger.error(f"❌ Text extraction failed for {file_format}: {str(e)}")
            return ""
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            text_content = []
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"⚠️ Failed to extract text from PDF page {page_num + 1}: {str(e)}")
                    continue
            
            return "\\n\\n".join(text_content)
            
        except Exception as e:
            logger.error(f"❌ PDF text extraction failed: {str(e)}")
            return ""
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return "\\n\\n".join(text_content)
            
        except Exception as e:
            logger.error(f"❌ DOCX text extraction failed: {str(e)}")
            return ""
    
    async def _extract_txt_text(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first, then fall back to latin-1
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                return file_content.decode('latin-1', errors='ignore')
                
        except Exception as e:
            logger.error(f"❌ TXT text extraction failed: {str(e)}")
            return ""
    
    async def _save_file(self, file_content: bytes, document_id: str, file_format: str) -> Path:
        """Save uploaded file to disk"""
        filename = f"{document_id}.{file_format}"
        file_path = self.upload_dir / filename
        
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        return file_path
    
    def _validate_document_relevance(self, text: str, filename: str) -> Dict[str, Any]:
        """Validate if document is relevant for threat modeling"""
        try:
            # Convert to lowercase for analysis
            content_lower = text.lower()
            filename_lower = filename.lower()
            
            # Define relevance indicators
            relevant_keywords = {
                'security': ['security', 'threat', 'vulnerability', 'risk', 'attack', 'encryption', 'authentication', 'authorization', 'firewall', 'malware', 'breach', 'compliance', 'audit'],
                'architecture': ['architecture', 'system', 'component', 'service', 'infrastructure', 'deployment', 'network', 'topology', 'design', 'platform'],
                'technology': ['api', 'database', 'server', 'cloud', 'aws', 'azure', 'gcp', 'kubernetes', 'docker', 'microservices', 'application'],
                'business': ['requirements', 'specification', 'business', 'process', 'workflow', 'data flow', 'integration', 'protocol']
            }
            
            # Irrelevant content indicators
            irrelevant_keywords = [
                'recipe', 'cooking', 'food', 'restaurant', 'menu',
                'personal', 'diary', 'journal', 'letter', 'email',
                'novel', 'story', 'fiction', 'poetry', 'literature',
                'shopping', 'grocery', 'purchase', 'buy', 'sale',
                'medical', 'health', 'doctor', 'patient', 'prescription',
                'travel', 'vacation', 'hotel', 'flight', 'tourism'
            ]
            
            # Count relevant keywords
            relevance_score = 0
            found_categories = []
            
            for category, keywords in relevant_keywords.items():
                category_matches = sum(1 for keyword in keywords if keyword in content_lower)
                if category_matches > 0:
                    relevance_score += category_matches
                    found_categories.append(category)
            
            # Check for irrelevant content
            irrelevant_matches = sum(1 for keyword in irrelevant_keywords if keyword in content_lower)
            
            # Check filename for technical indicators
            filename_relevant = any([
                'spec' in filename_lower, 'requirement' in filename_lower,
                'architecture' in filename_lower, 'design' in filename_lower,
                'security' in filename_lower, 'threat' in filename_lower,
                'system' in filename_lower, 'technical' in filename_lower
            ])
            
            # Calculate final relevance
            is_relevant = (
                relevance_score >= 3 or  # At least 3 relevant keywords
                filename_relevant or     # Filename suggests relevance
                (relevance_score > 0 and irrelevant_matches == 0)  # Some relevance and no irrelevant content
            )
            
            # Special case: very short documents are likely irrelevant
            if len(text.strip()) < 100:
                is_relevant = False
            
            return {
                'is_relevant': is_relevant,
                'relevance_score': relevance_score,
                'found_categories': found_categories,
                'irrelevant_matches': irrelevant_matches,
                'confidence': min(100, (relevance_score * 20) + (20 if filename_relevant else 0))
            }
            
        except Exception as e:
            logger.warning(f"Document relevance validation failed: {e}")
            # Default to relevant if validation fails
            return {
                'is_relevant': True,
                'relevance_score': 0,
                'found_categories': [],
                'irrelevant_matches': 0,
                'confidence': 50
            }
    
    def _analyze_document_content(self, text_content: str) -> Dict[str, Any]:
        """
        Analyze document content for threat modeling insights
        Extract key information that will help Claude provide better analysis
        """
        
        text_lower = text_content.lower()
        
        # Technology keywords
        technologies = []
        tech_keywords = {
            'cloud': ['aws', 'azure', 'gcp', 'google cloud', 'cloud', 'ec2', 's3', 'lambda'],
            'databases': ['mongodb', 'mysql', 'postgresql', 'redis', 'database', 'sql', 'nosql'],
            'web': ['api', 'rest', 'graphql', 'http', 'https', 'web service', 'microservice'],
            'auth': ['authentication', 'authorization', 'oauth', 'saml', 'jwt', 'token'],
            'containers': ['docker', 'kubernetes', 'container', 'pod', 'deployment'],
            'frameworks': ['react', 'node.js', 'express', 'spring', 'django', 'flask']
        }
        
        for category, keywords in tech_keywords.items():
            found_keywords = [kw for kw in keywords if kw in text_lower]
            if found_keywords:
                technologies.append({
                    "category": category,
                    "keywords": found_keywords
                })
        
        # Security keywords
        security_concerns = []
        security_keywords = [
            'security', 'privacy', 'encryption', 'ssl', 'tls', 'vulnerability', 
            'threat', 'risk', 'compliance', 'gdpr', 'pci', 'hipaa', 'audit'
        ]
        
        for keyword in security_keywords:
            if keyword in text_lower:
                security_concerns.append(keyword)
        
        # Data sensitivity indicators
        data_types = []
        data_keywords = [
            'personal data', 'pii', 'credit card', 'payment', 'user data', 
            'customer data', 'sensitive', 'confidential', 'financial'
        ]
        
        for keyword in data_keywords:
            if keyword in text_lower:
                data_types.append(keyword)
        
        return {
            "technologies": technologies,
            "security_concerns": security_concerns,
            "data_types": data_types,
            "has_architecture_info": any(word in text_lower for word in ['architecture', 'system', 'component', 'service']),
            "has_security_requirements": any(word in text_lower for word in ['requirement', 'must', 'shall', 'security', 'protect']),
            "complexity_score": min(len(technologies) + len(security_concerns), 10)  # 0-10 scale
        }

# Global instance
document_processor = DocumentProcessor()